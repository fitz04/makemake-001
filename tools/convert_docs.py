#!/usr/bin/env python3
"""
Project Nomad - 문서 변환 도구

Google Drive에서 다운로드한 .docx 파일들을 .txt로 변환합니다.
변환된 텍스트 파일은 AI 에이전트가 분석에 사용합니다.

사전 요구사항:
    pip install python-docx

사용법:
    python3 tools/convert_docs.py                          # 기본: docs/ -> docs_txt/
    python3 tools/convert_docs.py --input path/to/docx/    # 입력 디렉토리 지정
    python3 tools/convert_docs.py --output path/to/txt/    # 출력 디렉토리 지정
    python3 tools/convert_docs.py --single file.docx       # 단일 파일 변환
"""

import argparse
import sys
from pathlib import Path


def check_dependencies() -> bool:
    """python-docx 설치 여부 확인"""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        print("ERROR: python-docx가 설치되어 있지 않습니다.")
        print("설치: pip install python-docx")
        return False


def convert_docx_to_text(docx_path: Path) -> str:
    """
    .docx 파일을 텍스트로 변환합니다.

    테이블, 문단, 헤더를 포함하여 가능한 한 많은 구조를 보존합니다.
    """
    import docx

    doc = docx.Document(str(docx_path))
    lines: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # 문단 처리
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para and para.text.strip():
                # 스타일 기반 포맷팅
                style_name = para.style.name if para.style else ""
                text = para.text.strip()

                if "Heading 1" in style_name:
                    lines.append(f"\n{'=' * 60}")
                    lines.append(text)
                    lines.append("=" * 60)
                elif "Heading 2" in style_name:
                    lines.append(f"\n{'-' * 40}")
                    lines.append(text)
                    lines.append("-" * 40)
                elif "Heading 3" in style_name:
                    lines.append(f"\n### {text}")
                elif "List" in style_name or text.startswith(("- ", "* ", "• ")):
                    lines.append(f"  - {text.lstrip('- *• ')}")
                else:
                    lines.append(text)

        elif tag == "tbl":
            # 테이블 처리
            for table in doc.tables:
                if table._element is element:
                    lines.append("")
                    for row_idx, row in enumerate(table.rows):
                        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        line = " | ".join(cells)
                        lines.append(f"| {line} |")
                        if row_idx == 0:
                            separator = " | ".join(["---"] * len(cells))
                            lines.append(f"| {separator} |")
                    lines.append("")
                    break

    return "\n".join(lines)


def convert_single_file(input_path: Path, output_path: Path) -> bool:
    """단일 .docx 파일을 변환합니다."""
    try:
        print(f"  변환 중: {input_path.name}")
        text = convert_docx_to_text(input_path)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(text, encoding="utf-8")

        size_kb = output_path.stat().st_size / 1024
        print(f"    -> {output_path.name} ({size_kb:.1f} KB)")
        return True

    except Exception as e:
        print(f"  ERROR: {input_path.name} 변환 실패: {e}")
        return False


def convert_directory(input_dir: Path, output_dir: Path) -> tuple[int, int]:
    """
    디렉토리 내 모든 .docx 파일을 변환합니다.

    Returns:
        (성공 수, 실패 수)
    """
    docx_files = sorted(input_dir.glob("*.docx"))

    if not docx_files:
        print(f"WARNING: {input_dir}에서 .docx 파일을 찾을 수 없습니다.")
        return 0, 0

    print(f"발견된 .docx 파일: {len(docx_files)}개")
    print(f"출력 디렉토리: {output_dir}")
    print()

    output_dir.mkdir(parents=True, exist_ok=True)

    success = 0
    fail = 0

    for docx_path in docx_files:
        txt_name = docx_path.stem + ".txt"
        output_path = output_dir / txt_name

        if convert_single_file(docx_path, output_path):
            success += 1
        else:
            fail += 1

    return success, fail


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Project Nomad - .docx를 .txt로 변환",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "--input", "-i",
        type=Path,
        default=Path("docs"),
        help="입력 디렉토리 (기본: docs/)",
    )
    parser.add_argument(
        "--output", "-o",
        type=Path,
        default=Path("docs_txt"),
        help="출력 디렉토리 (기본: docs_txt/)",
    )
    parser.add_argument(
        "--single", "-s",
        type=Path,
        default=None,
        help="단일 파일 변환 모드",
    )
    args = parser.parse_args()

    if not check_dependencies():
        return 1

    print("=" * 50)
    print("Project Nomad - 문서 변환 도구")
    print("=" * 50)
    print()

    if args.single:
        # 단일 파일 모드
        if not args.single.exists():
            print(f"ERROR: 파일을 찾을 수 없습니다: {args.single}")
            return 1

        output_path = args.output / (args.single.stem + ".txt")
        ok = convert_single_file(args.single, output_path)
        return 0 if ok else 1

    else:
        # 디렉토리 모드
        if not args.input.exists():
            print(f"ERROR: 입력 디렉토리를 찾을 수 없습니다: {args.input}")
            return 1

        success, fail = convert_directory(args.input, args.output)

        print()
        print(f"완료: {success}개 성공, {fail}개 실패")
        return 0 if fail == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
